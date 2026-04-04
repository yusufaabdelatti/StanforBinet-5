import streamlit as st
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io, os, smtplib, re, json, base64
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from datetime import date
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np

# ReportLab
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import cm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                 TableStyle, Image as RLImage, HRFlowable,
                                 PageBreak, KeepTogether)

# ══════════════════════════════════════════════════════════════
#  CONFIG
# ══════════════════════════════════════════════════════════════
GMAIL_USER      = "Wijdan.psyc@gmail.com"
GMAIL_PASS      = "rias eeul lyuu stce"
RECIPIENT_EMAIL = "Wijdan.psyc@gmail.com"
LOGO_PATH       = os.path.join(os.path.dirname(__file__), "logo.png")

DEEP_BLUE     = "#1B3A6B"
MID_BLUE      = "#2E6DB4"
GOLD          = "#C8922A"
LIGHT_BG      = "#F5F7FA"

DEEP_BLUE_RGB = RGBColor(0x1B, 0x3A, 0x6B)
MID_BLUE_RGB  = RGBColor(0x2E, 0x6D, 0xB4)
GOLD_RGB      = RGBColor(0xC8, 0x92, 0x2A)
DARK_RGB      = RGBColor(0x1A, 0x1A, 0x2E)

# PDF palette
PDF_DEEP    = colors.HexColor('#1B3A6B')
PDF_MID     = colors.HexColor('#2E6DB4')
PDF_GOLD    = colors.HexColor('#C8922A')
PDF_LIGHT   = colors.HexColor('#F5F7FA')
PDF_BORDER  = colors.HexColor('#C5D3F5')
PDF_CREAM   = colors.HexColor('#EEF2FF')

# ══════════════════════════════════════════════════════════════
#  SB-5 STRUCTURE
# ══════════════════════════════════════════════════════════════
IQ_SCORES     = ["FSIQ","NVIQ","VIQ"]
FACTOR_SCORES = ["FR","KN","QR","VS","WM"]

IQ_LABELS = {
    "FSIQ": {"en":"Full Scale IQ",          "ar":"درجة الذكاء الكلية"},
    "NVIQ": {"en":"Nonverbal IQ",           "ar":"درجة المجال غير اللفظي"},
    "VIQ":  {"en":"Verbal IQ",              "ar":"درجة المجال اللفظي"},
}
FACTOR_LABELS = {
    "FR": {"en":"Fluid Reasoning",          "ar":"الاستدلال التحليلي"},
    "KN": {"en":"Knowledge",               "ar":"المعلومات"},
    "QR": {"en":"Quantitative Reasoning",  "ar":"الاستدلال الكمي"},
    "VS": {"en":"Visual-Spatial",          "ar":"المعالجة البصرية المكانية"},
    "WM": {"en":"Working Memory",          "ar":"الذاكرة العاملة"},
}

CLASSIFICATIONS = [
    (145,160,"Very Gifted",         "موهوب بشدة",        "#1565C0"),
    (130,144,"Gifted",              "موهوب",             "#1976D2"),
    (120,129,"Superior",            "متفوق",             "#0288D1"),
    (110,119,"High Average",        "متوسط مرتفع",       "#00897B"),
    (90, 109,"Average",             "متوسط",             "#388E3C"),
    (80,  89,"Low Average",         "متوسط منخفض",       "#F9A825"),
    (70,  79,"Borderline",          "الفئة البينية",     "#EF6C00"),
    (55,  69,"Mild Impairment",     "إعاقة بسيطة",      "#E53935"),
    (40,  54,"Moderate Impairment", "إعاقة متوسطة",     "#B71C1C"),
    (0,   39,"Severe Impairment",   "إعاقة شديدة",      "#880E4F"),
]

def classify(ss):
    if ss is None: return "Unknown","غير محدد","#888888"
    for lo,hi,en,ar,col in CLASSIFICATIONS:
        if lo <= ss <= hi: return en,ar,col
    return "Unknown","غير محدد","#888888"

def percentile_from_ss(ss):
    if ss is None: return 0
    from math import erf,sqrt
    z = (ss-100)/15.0
    p = 0.5*(1+erf(z/sqrt(2)))
    return max(1,min(99,round(p*100)))

# ══════════════════════════════════════════════════════════════
#  STEP 1: EXTRACT TEXT FROM UPLOADED WORD DOCUMENT
# ══════════════════════════════════════════════════════════════
def extract_docx_text(uploaded_file) -> str:
    """Extract all text from the uploaded Word (.docx) file using python-docx."""
    uploaded_file.seek(0)
    doc = Document(io.BytesIO(uploaded_file.read()))
    parts = []
    # Paragraphs
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            parts.append(t)
    # Tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)

# ══════════════════════════════════════════════════════════════
#  STEP 2: PARSE DATA FROM RAW TEXT VIA GROQ
# ══════════════════════════════════════════════════════════════
def extract_data_from_text(raw_text: str) -> dict:
    """Send raw PDF text to Groq and get structured JSON data back."""
    prompt = f"""You are a clinical psychologist assistant. The following is raw text extracted from an Arabic Stanford-Binet 5 (SB5) psychological assessment report.

Extract ALL available information and return ONLY a valid JSON object — no explanation, no markdown, no backticks.

Extract these fields (use null for any missing value):
{{
  "name": "examinee full name",
  "dob": "date of birth",
  "age": "age as stated",
  "gender": "male or female (translate from Arabic if needed)",
  "grade": "school grade or level",
  "school": "school or institution name",
  "examiner": "examiner name",
  "test_date": "date of testing",
  "referral": "referral source",
  "complaints": "reason for referral / presenting complaints",
  "behavioral_obs": "behavioral observations during testing",
  "background": "background information / history",
  "FSIQ": integer_or_null,
  "FSIQ_ci": "confidence interval as string e.g. 91-99 or null",
  "NVIQ": integer_or_null,
  "NVIQ_ci": "string or null",
  "VIQ": integer_or_null,
  "VIQ_ci": "string or null",
  "FR": integer_or_null,
  "FR_ci": "string or null",
  "KN": integer_or_null,
  "KN_ci": "string or null",
  "QR": integer_or_null,
  "QR_ci": "string or null",
  "VS": integer_or_null,
  "VS_ci": "string or null",
  "WM": integer_or_null,
  "WM_ci": "string or null",
  "nv_fr": integer_or_null,
  "v_fr": integer_or_null,
  "nv_kn": integer_or_null,
  "v_kn": integer_or_null,
  "nv_qr": integer_or_null,
  "v_qr": integer_or_null,
  "nv_vs": integer_or_null,
  "v_vs": integer_or_null,
  "nv_wm": integer_or_null,
  "v_wm": integer_or_null
}}

Notes:
- درجة الذكاء الكلية = FSIQ
- المجال غير اللفظي / غير لفظي = NVIQ
- المجال اللفظي / لفظي = VIQ
- الاستدلال التحليلي / الاستدلال الطلاقة = FR
- المعلومات / المعرفة = KN
- الاستدلال الكمي = QR
- المعالجة البصرية المكانية / البصري المكاني = VS
- الذاكرة العاملة = WM
- غير لفظي = NV (nonverbal subtest)
- لفظي = V (verbal subtest)
- Any number in parentheses after a score is likely the confidence interval
- Percentile ranks (رتبة مئينية) should NOT be confused with standard scores
- Standard scores for IQ are typically 40-160; subtest scaled scores are typically 1-19

RAW REPORT TEXT:
{raw_text[:8000]}
"""
    client = Groq(api_key=st.secrets["GROQ_API_KEY"])
    r = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role":"user","content":prompt}],
        max_tokens=1500,
        temperature=0.1,
    )
    raw_json = r.choices[0].message.content.strip()
    # Strip any accidental markdown fences
    raw_json = re.sub(r"^```[a-z]*\n?","",raw_json)
    raw_json = re.sub(r"\n?```$","",raw_json)
    try:
        return json.loads(raw_json)
    except Exception:
        # Try to extract JSON substring
        m = re.search(r'\{.*\}', raw_json, re.DOTALL)
        if m:
            return json.loads(m.group())
        return {}

# ══════════════════════════════════════════════════════════════
#  CHARTS
# ══════════════════════════════════════════════════════════════
def make_profile_chart(data: dict) -> bytes:
    all_keys   = ["FSIQ","NVIQ","VIQ","FR","KN","QR","VS","WM"]
    all_labels, all_vals, all_colors = [], [], []
    for k in all_keys:
        v = data.get(k)
        if v is None: continue
        lbl = (IQ_LABELS if k in IQ_LABELS else FACTOR_LABELS)[k]["en"]
        _,_,col = classify(v)
        all_labels.append(lbl); all_vals.append(v); all_colors.append(col)
    if not all_vals: return None

    fig, ax = plt.subplots(figsize=(12,7))
    fig.patch.set_facecolor('#FAFBFF'); ax.set_facecolor('#FAFBFF')
    y_pos = np.arange(len(all_labels))

    bands = [(40,54,"#FCE4EC"),(55,69,"#FFEBEE"),(70,79,"#FFF3E0"),
             (80,89,"#FFFDE7"),(90,109,"#F1F8E9"),(110,119,"#E3F2FD"),
             (120,129,"#E1F5FE"),(130,160,"#E8EAF6")]
    for lo,hi,col in bands:
        ax.axvspan(lo,hi,alpha=0.3,color=col,zorder=0)

    bars = ax.barh(y_pos, all_vals, color=all_colors, height=0.6,
                   edgecolor='white', linewidth=1.2, zorder=3)
    n_iq = sum(1 for k in all_keys if k in IQ_LABELS and data.get(k))
    if 0 < n_iq < len(all_labels):
        ax.axhline(y=n_iq-0.5, color='#BDBDBD', linestyle='--', linewidth=1.2, zorder=4)

    ax.axvline(x=100, color='#1B3A6B', linestyle='-', linewidth=2, alpha=0.6, label='Mean (100)')
    ax.axvline(x=85,  color='#EF6C00', linestyle=':', linewidth=1.2, alpha=0.5)
    ax.axvline(x=115, color='#EF6C00', linestyle=':', linewidth=1.2, alpha=0.5)

    for bar_, val in zip(bars, all_vals):
        en_c,_,_ = classify(val)
        pct = percentile_from_ss(val)
        label = f"{val}  ({en_c}, {pct}th %ile)"
        ax.text(bar_.get_width()+1, bar_.get_y()+bar_.get_height()/2,
                label, va='center', ha='left', fontsize=8.5, color='#1A1A2E')

    ax.set_yticks(y_pos); ax.set_yticklabels(all_labels, fontsize=10.5)
    ax.set_xlim(40,175)
    ax.set_xlabel("Standard Score  (Mean=100, SD=15)", fontsize=11, color='#1A1A2E')
    ax.set_title("Stanford-Binet 5 — Full Score Profile", fontsize=13,
                 fontweight='bold', color='#1B3A6B', pad=12)
    ax.legend(fontsize=9, framealpha=0.7)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    ax.grid(axis='x', linestyle=':', alpha=0.4, zorder=1)
    plt.tight_layout()
    buf=io.BytesIO(); plt.savefig(buf,format='png',dpi=150,bbox_inches='tight')
    plt.close(fig); buf.seek(0); return buf.read()

def make_factor_radar(data: dict) -> bytes:
    labels = [FACTOR_LABELS[k]["en"] for k in FACTOR_SCORES if data.get(k)]
    vals   = [data[k] for k in FACTOR_SCORES if data.get(k)]
    if len(vals) < 3: return None
    N = len(labels)
    angles = np.linspace(0,2*np.pi,N,endpoint=False).tolist()
    vals_  = vals+[vals[0]]; angles_= angles+[angles[0]]
    fig,ax = plt.subplots(figsize=(6,6),subplot_kw=dict(polar=True))
    fig.patch.set_facecolor('#FAFBFF'); ax.set_facecolor('#EEF2FF')
    for ref,col in [(70,'#EF6C00'),(100,'#1B3A6B'),(130,'#0288D1')]:
        ref_n = ref/160.0
        ax.plot(angles_,[ref_n]*len(angles_),'--',color=col,linewidth=0.9,alpha=0.5)
    vals_n = [v/160.0 for v in vals_]
    ax.plot(angles_,vals_n,'o-',linewidth=2.2,color='#2E6DB4',markersize=7)
    ax.fill(angles_,vals_n,alpha=0.28,color='#2E6DB4')
    ax.set_xticks(angles[:-1] if len(angles)==N+1 else angles)
    ax.set_xticklabels(labels,size=9.5)
    ax.set_ylim(0,1)
    ax.set_yticks([70/160,85/160,100/160,115/160,130/160])
    ax.set_yticklabels(['70','85','100','115','130'],size=7.5,color='#666')
    ax.set_title("Factor Index Radar Profile",size=12,fontweight='bold',color='#1B3A6B',pad=18)
    plt.tight_layout()
    buf=io.BytesIO(); plt.savefig(buf,format='png',dpi=150,bbox_inches='tight')
    plt.close(fig); buf.seek(0); return buf.read()

def make_classification_gauge(fsiq: int) -> bytes:
    fig,ax = plt.subplots(figsize=(10,4))
    fig.patch.set_facecolor('#FAFBFF'); ax.set_facecolor('#FAFBFF')
    bands_ordered = [
        (40, 54, "#B71C1C","Moderate\nImpairment"),
        (55, 69, "#E53935","Mild\nImpairment"),
        (70, 79, "#EF6C00","Borderline"),
        (80, 89, "#F9A825","Low\nAverage"),
        (90,109, "#388E3C","Average"),
        (110,119,"#00897B","High\nAverage"),
        (120,129,"#0288D1","Superior"),
        (130,145,"#1565C0","Gifted+"),
    ]
    for lo,hi,col,lbl in bands_ordered:
        ax.barh(0,hi-lo,left=lo,height=0.55,color=col,alpha=0.85,edgecolor='white',linewidth=1.5)
        ax.text((lo+hi)/2,0,lbl,ha='center',va='center',fontsize=7.5,
                color='white',fontweight='bold')
    ax.annotate('',xy=(fsiq,0.35),xytext=(fsiq,0.72),
                arrowprops=dict(arrowstyle='->',color='#1A1A2E',lw=2.8))
    en_c,_,_ = classify(fsiq)
    pct = percentile_from_ss(fsiq)
    ax.text(fsiq,0.88,f"FSIQ = {fsiq}\n{en_c} | {pct}th percentile",
            ha='center',va='bottom',fontsize=10.5,fontweight='bold',color='#1A1A2E')
    ax.set_xlim(40,145); ax.set_ylim(-0.45,1.15)
    ax.set_xlabel("Standard Score",fontsize=10)
    ax.set_yticks([])
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    ax.spines['left'].set_visible(False)
    ax.set_title("Full Scale IQ — Classification Gauge",fontsize=12,
                 fontweight='bold',color='#1B3A6B',pad=8)
    plt.tight_layout()
    buf=io.BytesIO(); plt.savefig(buf,format='png',dpi=150,bbox_inches='tight')
    plt.close(fig); buf.seek(0); return buf.read()

def make_subtest_chart(data: dict) -> bytes:
    keys = FACTOR_SCORES
    nv_vals = [data.get(f"nv_{k.lower()}") for k in keys]
    v_vals  = [data.get(f"v_{k.lower()}")  for k in keys]
    if all(v is None for v in nv_vals+v_vals): return None
    labels = [FACTOR_LABELS[k]["en"] for k in keys]
    x = np.arange(len(labels)); w = 0.35
    fig,ax = plt.subplots(figsize=(11,5.5))
    fig.patch.set_facecolor('#FAFBFF'); ax.set_facecolor('#FAFBFF')
    nv_c = [v if v is not None else 0 for v in nv_vals]
    v_c  = [v if v is not None else 0 for v in v_vals]
    b1=ax.bar(x-w/2,nv_c,w,label='Nonverbal',color='#2E6DB4',edgecolor='white',linewidth=0.8)
    b2=ax.bar(x+w/2,v_c, w,label='Verbal',   color='#C8922A',edgecolor='white',linewidth=0.8)
    ax.axhline(y=10,color='#1B3A6B',linestyle='--',linewidth=1.8,alpha=0.6,label='Mean (10)')
    ax.axhline(y=7, color='#EF6C00',linestyle=':',linewidth=1.2,alpha=0.5)
    ax.axhline(y=13,color='#EF6C00',linestyle=':',linewidth=1.2,alpha=0.5)
    for bar_ in list(b1)+list(b2):
        h=bar_.get_height()
        if h>0: ax.text(bar_.get_x()+bar_.get_width()/2.,h+0.15,str(int(h)),
                        ha='center',va='bottom',fontsize=9.5,fontweight='bold')
    ax.set_xticks(x); ax.set_xticklabels(labels,rotation=15,ha='right',fontsize=10)
    ax.set_ylabel("Scaled Score  (Mean=10, SD=3)",fontsize=10)
    ax.set_title("Subtest Scaled Scores — Nonverbal vs. Verbal",fontsize=12,
                 fontweight='bold',color='#1B3A6B')
    ax.set_ylim(0,21); ax.legend(fontsize=10)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    ax.grid(axis='y',linestyle=':',alpha=0.4)
    plt.tight_layout()
    buf=io.BytesIO(); plt.savefig(buf,format='png',dpi=150,bbox_inches='tight')
    plt.close(fig); buf.seek(0); return buf.read()

def make_discrepancy_chart(data: dict) -> bytes:
    """Bar chart comparing NV vs V IQ and showing discrepancy significance."""
    nviq = data.get("NVIQ"); viq = data.get("VIQ")
    if not nviq or not viq: return None
    disc = abs(nviq - viq)
    fig, axes = plt.subplots(1, 2, figsize=(11, 4.5))
    fig.patch.set_facecolor('#FAFBFF')
    # Left: NV vs V comparison
    ax = axes[0]; ax.set_facecolor('#FAFBFF')
    cats = ["Nonverbal IQ\n(NVIQ)", "Verbal IQ\n(VIQ)"]
    vals = [nviq, viq]
    cols = ["#2E6DB4","#C8922A"]
    bars = ax.bar(cats, vals, color=cols, width=0.45, edgecolor='white', linewidth=1.5)
    ax.axhline(y=100, color='#1B3A6B', linestyle='--', linewidth=1.5, alpha=0.6)
    for bar_,val in zip(bars,vals):
        en_c,_,_ = classify(val)
        ax.text(bar_.get_x()+bar_.get_width()/2., val+1.5,
                f"{val}\n({en_c})", ha='center', va='bottom', fontsize=10, fontweight='bold')
    ax.set_ylim(50, max(vals)+25)
    ax.set_title("NV vs. V IQ Comparison", fontsize=11, fontweight='bold', color='#1B3A6B')
    ax.set_ylabel("Standard Score", fontsize=10)
    ax.spines['top'].set_visible(False); ax.spines['right'].set_visible(False)
    # Right: discrepancy significance
    ax2 = axes[1]; ax2.set_facecolor('#FAFBFF')
    sig_color = "#E53935" if disc >= 23 else "#F9A825" if disc >= 15 else "#388E3C"
    sig_label = "Statistically Significant\n(p < .05)" if disc >= 15 else "Not Significant"
    ax2.barh(0, disc, height=0.4, color=sig_color, edgecolor='white')
    ax2.axvline(x=15, color='#F9A825', linestyle='--', linewidth=1.5, label='Significant (15)')
    ax2.axvline(x=23, color='#E53935', linestyle='--', linewidth=1.5, label='Highly Sig. (23)')
    ax2.text(disc+0.5, 0, f"Discrepancy = {disc}\n{sig_label}",
             va='center', ha='left', fontsize=10, fontweight='bold', color='#1A1A2E')
    ax2.set_xlim(0, max(30, disc+15))
    ax2.set_yticks([])
    ax2.set_xlabel("Point Difference (|NVIQ − VIQ|)", fontsize=10)
    ax2.set_title("NV–V Discrepancy Analysis", fontsize=11, fontweight='bold', color='#1B3A6B')
    ax2.legend(fontsize=9, framealpha=0.7)
    ax2.spines['top'].set_visible(False); ax2.spines['right'].set_visible(False)
    plt.tight_layout()
    buf=io.BytesIO(); plt.savefig(buf,format='png',dpi=150,bbox_inches='tight')
    plt.close(fig); buf.seek(0); return buf.read()

# ══════════════════════════════════════════════════════════════
#  STEP 3: GROQ REPORT GENERATION
# ══════════════════════════════════════════════════════════════
def build_score_summary(data: dict) -> str:
    lines = ["IQ SCORES:"]
    for k in IQ_SCORES:
        v = data.get(k)
        if v:
            en_c,_,_ = classify(v); pct=percentile_from_ss(v); ci=data.get(f"{k}_ci","—")
            lines.append(f"  {IQ_LABELS[k]['en']}: SS={v}, %ile={pct}th, 90%CI={ci}, Classification={en_c}")
    lines.append("\nFACTOR INDEX SCORES:")
    for k in FACTOR_SCORES:
        v = data.get(k)
        if v:
            en_c,_,_ = classify(v); pct=percentile_from_ss(v); ci=data.get(f"{k}_ci","—")
            lines.append(f"  {FACTOR_LABELS[k]['en']}: SS={v}, %ile={pct}th, 90%CI={ci}, Classification={en_c}")
    lines.append("\nSUBTEST SCALED SCORES (NV | V):")
    for k in FACTOR_SCORES:
        nv=data.get(f"nv_{k.lower()}"); v=data.get(f"v_{k.lower()}")
        if nv or v:
            lines.append(f"  {FACTOR_LABELS[k]['en']}: NV={nv or '—'} | V={v or '—'}")
    return "\n".join(lines)

def generate_en_report(data: dict) -> str:
    score_summary = build_score_summary(data)
    name     = data.get("name","the examinee") or "the examinee"
    age      = data.get("age","—") or "—"
    gender   = data.get("gender","—") or "—"
    pronoun  = "he" if "male" in str(gender).lower() else "she"
    examiner = data.get("examiner","—") or "—"
    referral = data.get("referral","—") or "—"
    complaints = data.get("complaints","Not provided") or "Not provided"
    behavioral = data.get("behavioral_obs","Not provided") or "Not provided"
    background = data.get("background","Not provided") or "Not provided"

    prompt = f"""You are a senior licensed psychologist writing a world-class Stanford-Binet 5 (SB5) psychological assessment report.
Use the most current research on SB5 (Roid, 2003; Roid et al., 2016) and CHC theory of intelligence.

EXAMINEE: {name} | AGE: {age} | GENDER: {gender} | EXAMINER: {examiner}
REFERRAL SOURCE: {referral}
REASON FOR REFERRAL: {complaints}
TEST DATE: {data.get("test_date","—") or "—"}
REPORT DATE: {date.today().strftime('%B %d, %Y')}

SCORE SUMMARY:
{score_summary}

BEHAVIORAL OBSERVATIONS: {behavioral}
BACKGROUND INFORMATION: {background}

Write a COMPREHENSIVE, PREMIUM SB5 report. Use formal psychoeducational language. Be very specific to the actual scores.
No markdown symbols (**, ##, ---, etc.). Section titles: ALL CAPS on their own line.
Use {pronoun}/{pronoun}s consistently. Reference specific scores and percentiles throughout.

REPORT STRUCTURE:

STANFORD-BINET INTELLIGENCE SCALES, FIFTH EDITION — PSYCHOLOGICAL REPORT
Name | {name}
Date of Birth | {data.get("dob","—") or "—"}
Age | {age}
Gender | {gender}
Examiner | {examiner}
Test Date | {data.get("test_date","—") or "—"}
Report Date | {date.today().strftime('%B %d, %Y')}
Referral Source | {referral}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

REASON FOR REFERRAL
3–5 sentences: who referred, what was the clinical question, what this evaluation addresses.

BACKGROUND INFORMATION
Relevant developmental, educational, medical, and family history from the provided background.

TESTS ADMINISTERED
List SB5 and any other assessments mentioned.

BEHAVIORAL OBSERVATIONS
Describe behavior during testing: cooperation, attention, affect, language, motor, approach.
Note any factors that may affect validity.

ASSESSMENT RESULTS AND INTERPRETATION

1. FULL SCALE IQ (FSIQ)
Interpret FSIQ deeply: what it represents, confidence interval, percentile, classification, clinical meaning.
Reference CHC theory (Gf-Gc framework).

2. NONVERBAL IQ (NVIQ) AND VERBAL IQ (VIQ)
Compare NV and V domains. Discuss discrepancy (statistically significant at p<.05 if |NV-V|≥15).
Clinical implications of any discrepancy.

3. FLUID REASONING (FR)
Interpret factor score, NV and V subtest scores, inductive/deductive reasoning abilities.

4. KNOWLEDGE (KN)
Interpret crystallized intelligence, vocabulary, acquired knowledge.

5. QUANTITATIVE REASONING (QR)
Interpret numerical and mathematical reasoning abilities.

6. VISUAL-SPATIAL PROCESSING (VS)
Interpret spatial visualization, pattern recognition, visuoconstructive abilities.

7. WORKING MEMORY (WM)
Interpret short-term memory, attention, cognitive control.

8. STRENGTHS AND WEAKNESSES PROFILE
Identify relative strengths (highest scores) and weaknesses (lowest scores).
Discuss intra-cognitive variability. Is the profile consistent or scattered?

9. DIAGNOSTIC IMPRESSIONS
Patterns that emerge. Reference relevant diagnostic considerations as hypotheses only.
Do NOT provide a formal diagnosis.

10. RECOMMENDATIONS
Provide 10–12 specific, evidence-based recommendations across:
a) Educational accommodations and classroom strategies
b) Intervention priorities (cognitive, academic, behavioral)
c) Further evaluation needs
d) Family and home support strategies
e) Therapeutic or clinical referrals if indicated

11. SUMMARY
A concise 2-paragraph executive summary for school teams and specialists.

PARENT-FRIENDLY SUMMARY
Write a plain-language 2-paragraph explanation for parents/caregivers with NO jargon.
Explain what the test found, what it means for their child day-to-day, and the 3 most important next steps.
"""
    client = Groq(api_key=st.secrets["GROQ_API_KEY"])
    r = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role":"user","content":prompt}],
        max_tokens=4500,
    )
    return r.choices[0].message.content.strip()

def generate_ar_report(data: dict) -> str:
    name     = data.get("name","المفحوص") or "المفحوص"
    age      = data.get("age","—") or "—"
    gender   = data.get("gender","—") or "—"
    examiner = data.get("examiner","—") or "—"
    referral = data.get("referral","—") or "—"

    ar_scores = []
    for k in IQ_SCORES:
        v = data.get(k)
        if v:
            _,ar_c,_ = classify(v); pct=percentile_from_ss(v); ci=data.get(f"{k}_ci","—")
            ar_scores.append(f"  {IQ_LABELS[k]['ar']}: {v} — {ar_c} (مئيني: {pct}، مدى ثقة 90%: {ci})")
    for k in FACTOR_SCORES:
        v = data.get(k)
        if v:
            _,ar_c,_ = classify(v); pct=percentile_from_ss(v); ci=data.get(f"{k}_ci","—")
            ar_scores.append(f"  {FACTOR_LABELS[k]['ar']}: {v} — {ar_c} (مئيني: {pct}، مدى ثقة: {ci})")

    prompt = f"""أنت طبيب نفسي متخصص تكتب تقريراً سريرياً شاملاً ومتقدماً لمقياس ستانفورد-بينيه الصورة الخامسة (SB5).
استخدم نظرية CHC في الذكاء (Cattell-Horn-Carroll) وأحدث الأبحاث المتعلقة بالمقياس.

المفحوص: {name} | العمر: {age} | النوع: {gender} | الفاحص: {examiner}
جهة الإحالة: {referral}
سبب الإحالة: {data.get("complaints","—") or "—"}
تاريخ التطبيق: {data.get("test_date","—") or "—"}
تاريخ التقرير: {date.today().strftime('%Y/%m/%d')}

ملخص الدرجات:
{chr(10).join(ar_scores)}

الملاحظات السلوكية: {data.get("behavioral_obs","لم تُذكر") or "لم تُذكر"}
المعلومات الأساسية: {data.get("background","لم تُذكر") or "لم تُذكر"}

اكتب تقريراً سريرياً نفسياً شاملاً ومتقدماً بالعربية الفصحى.
لا تستخدم رموز markdown (**, ##, ---, إلخ). عناوين الأقسام: أرقام + عناوين واضحة.
لا إنجليزية إلا للاختصارات المقبولة (SB5, IQ, CHC, FSIQ, NVIQ, VIQ).

تقرير مقياس ستانفورد-بينيه للذكاء — الصورة الخامسة
الاسم | {name}
تاريخ الميلاد | {data.get("dob","—") or "—"}
العمر | {age}
النوع | {gender}
الفاحص | {examiner}
تاريخ التطبيق | {data.get("test_date","—") or "—"}
تاريخ التقرير | {date.today().strftime('%Y/%m/%d')}
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━

١. سبب الإحالة
فقرة من ٣-٥ جمل: من أحال وما هو التساؤل الإكلينيكي.

٢. المعلومات الأساسية
التاريخ التطوري والتعليمي والطبي والأسري ذات الصلة.

٣. الاختبارات المطبقة
اذكر المقياس وأي أدوات أخرى ذُكرت.

٤. الملاحظات السلوكية
وصف سلوك المفحوص أثناء التطبيق. اذكر العوامل المؤثرة في صدق النتائج.

٥. نتائج التقييم وتفسيرها

أ. درجة الذكاء الكلية (FSIQ)
تفسير معمق: الدرجة، فترة الثقة، المئيني، الفئة، الدلالة الإكلينيكية في إطار نظرية CHC.

ب. الذكاء غير اللفظي (NVIQ) والذكاء اللفظي (VIQ)
مقارنة المجالين. الدلالة الإكلينيكية لأي فرق ذي دلالة إحصائية.

ج. الاستدلال التحليلي (FR)
تفسير عامل الاستدلال السائل.

د. المعلومات (KN)
تفسير الذكاء المتبلور والمعرفة العامة.

هـ. الاستدلال الكمي (QR)
القدرة الرياضية والعددية.

و. المعالجة البصرية المكانية (VS)
القدرة على التصور والمعالجة البصرية.

ز. الذاكرة العاملة (WM)
الذاكرة قصيرة المدى، الانتباه، الكفاءة المعرفية.

٦. نقاط القوة والقصور النسبي
تحديد نقاط القوة والقصور النسبي بناءً على الملف الكامل.

٧. الانطباعات التشخيصية
الأنماط الإكلينيكية الظاهرة. لا تقديم تشخيص رسمي — فرضيات إكلينيكية فقط.

٨. التوصيات
قدم ١٠-١٢ توصية محددة ومبنية على الأدلة في المجالات:
أ) التعليمية والأكاديمية   ب) التدخل العلاجي   ج) تقييمات إضافية   د) دعم الأسرة   هـ) الإحالات

٩. الملخص
فقرتان موجزتان للفريق المتخصص.

ملخص للوالدين
فقرتان بلغة مبسطة للأسرة بدون مصطلحات تخصصية.
اشرح ما وجده الاختبار وماذا يعني لطفلهم في حياته اليومية وأهم ٣ خطوات للأمام.
"""
    client = Groq(api_key=st.secrets["GROQ_API_KEY"])
    r = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role":"user","content":prompt}],
        max_tokens=4500,
    )
    return r.choices[0].message.content.strip()

# ══════════════════════════════════════════════════════════════
#  STEP 4A: BUILD ENGLISH PDF REPORT
# ══════════════════════════════════════════════════════════════
def _pdf_styles():
    S = {}
    S['title']    = ParagraphStyle('title',    fontName='Helvetica-Bold',  fontSize=17, textColor=PDF_DEEP,  spaceAfter=3,  alignment=TA_CENTER)
    S['subtitle'] = ParagraphStyle('subtitle', fontName='Helvetica',       fontSize=9,  textColor=PDF_GOLD,  spaceAfter=2,  alignment=TA_CENTER)
    S['section']  = ParagraphStyle('section',  fontName='Helvetica-Bold',  fontSize=11, textColor=PDF_DEEP,  spaceBefore=14,spaceAfter=4)
    S['body']     = ParagraphStyle('body',     fontName='Helvetica',       fontSize=9.5,textColor=colors.HexColor('#1A1A2E'),leading=14,spaceAfter=5)
    S['small']    = ParagraphStyle('small',    fontName='Helvetica',       fontSize=8,  textColor=PDF_MID,   leading=11)
    S['bold']     = ParagraphStyle('bold',     fontName='Helvetica-Bold',  fontSize=9.5,textColor=colors.HexColor('#1A1A2E'),leading=14,spaceAfter=3)
    S['parent']   = ParagraphStyle('parent',   fontName='Helvetica',       fontSize=10.5,textColor=colors.HexColor('#1A1A2E'),leading=16,spaceAfter=5,leftIndent=8)
    return S

def _hdr_style(): return colors.HexColor('#1B3A6B')
def _shade1():    return colors.HexColor('#EEF2FF')
def _shade2():    return colors.HexColor('#FFF8EE')

def _score_tbl(rows, col_widths, header_cmds=None):
    t = Table(rows, colWidths=col_widths, repeatRows=1)
    cmds = [
        ('BACKGROUND',(0,0),(-1,0), _hdr_style()),
        ('TEXTCOLOR', (0,0),(-1,0), colors.white),
        ('FONTNAME',  (0,0),(-1,0),'Helvetica-Bold'),
        ('FONTSIZE',  (0,0),(-1,-1),8.5),
        ('BOX',       (0,0),(-1,-1),0.5,PDF_BORDER),
        ('INNERGRID', (0,0),(-1,-1),0.3,PDF_BORDER),
        ('VALIGN',    (0,0),(-1,-1),'MIDDLE'),
        ('TOPPADDING',(0,0),(-1,-1),4),
        ('BOTTOMPADDING',(0,0),(-1,-1),4),
        ('LEFTPADDING',(0,0),(-1,-1),6),
    ]
    if header_cmds: cmds.extend(header_cmds)
    t.setStyle(TableStyle(cmds))
    return t

def build_pdf_report(report_text, data, charts, center_name="", logo_bytes=None) -> io.BytesIO:
    buf = io.BytesIO()
    W_page, H_page = A4
    doc = SimpleDocTemplate(buf, pagesize=A4,
        leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
    W = W_page - 4*cm
    S = _pdf_styles()
    story = []

    # ── Logo + Center Name + Title ──
    # Uploaded logo takes priority, fallback to file logo
    logo_shown = False
    if logo_bytes:
        try:
            logo = RLImage(io.BytesIO(logo_bytes), width=4.5*cm, height=2*cm)
            logo.hAlign = 'CENTER'
            story.append(logo)
            story.append(Spacer(1, 4))
            logo_shown = True
        except: pass
    if not logo_shown and os.path.exists(LOGO_PATH):
        try:
            logo = RLImage(LOGO_PATH, width=4.5*cm, height=2*cm)
            logo.hAlign = 'CENTER'
            story.append(logo)
            story.append(Spacer(1, 4))
        except: pass

    if center_name:
        story.append(Paragraph(center_name, ParagraphStyle('center_name',
            fontName='Helvetica-Bold', fontSize=13, textColor=PDF_DEEP,
            spaceAfter=4, alignment=TA_CENTER)))

    story.append(Paragraph("Stanford-Binet Intelligence Scales, Fifth Edition", S['title']))
    story.append(Paragraph("Psychological Assessment Report  ·  SB5 · Roid (2003)", S['subtitle']))
    story.append(HRFlowable(width=W, thickness=2, color=PDF_GOLD, spaceAfter=10))

    # ── Demographics table ──
    name    = data.get("name","—") or "—"
    age     = data.get("age","—") or "—"
    gender  = data.get("gender","—") or "—"
    examiner= data.get("examiner","—") or "—"
    dob     = data.get("dob","—") or "—"
    tdate   = data.get("test_date","—") or "—"
    ref     = data.get("referral","—") or "—"

    demo_rows = [
        [Paragraph('<b>Name</b>',S['small']),    Paragraph(name,S['body']),
         Paragraph('<b>Date of Birth</b>',S['small']), Paragraph(dob,S['body'])],
        [Paragraph('<b>Age</b>',S['small']),     Paragraph(age,S['body']),
         Paragraph('<b>Gender</b>',S['small']),  Paragraph(gender,S['body'])],
        [Paragraph('<b>Examiner</b>',S['small']),Paragraph(examiner,S['body']),
         Paragraph('<b>Test Date</b>',S['small']),Paragraph(tdate,S['body'])],
        [Paragraph('<b>Referral</b>',S['small']),Paragraph(ref,S['body']),
         Paragraph('<b>Report Date</b>',S['small']),Paragraph(date.today().strftime('%B %d, %Y'),S['body'])],
    ]
    demo_tbl = Table(demo_rows, colWidths=[2.5*cm,6.2*cm,3*cm,5.8*cm])
    demo_tbl.setStyle(TableStyle([
        ('BACKGROUND',(0,0),(-1,-1),_shade1()),
        ('BOX',(0,0),(-1,-1),0.5,PDF_BORDER),
        ('INNERGRID',(0,0),(-1,-1),0.3,PDF_BORDER),
        ('VALIGN',(0,0),(-1,-1),'MIDDLE'),
        ('TOPPADDING',(0,0),(-1,-1),5),
        ('BOTTOMPADDING',(0,0),(-1,-1),5),
        ('LEFTPADDING',(0,0),(-1,-1),6),
    ]))
    story.append(KeepTogether([demo_tbl, Spacer(1,10)]))

    # ── IQ + Factor Score Summary Table ──
    story.append(KeepTogether([
        Paragraph("IQ AND FACTOR INDEX SCORE SUMMARY", S['section']),
        _build_score_table(data, S, W),
        Spacer(1,8)
    ]))

    # ── Subtest Table ──
    sub_tbl = _build_subtest_table(data, S, W)
    if sub_tbl:
        story.append(KeepTogether([
            Paragraph("SUBTEST SCALED SCORES", S['section']),
            sub_tbl,
            Spacer(1,8)
        ]))

    # ── Classification Colour Legend ──
    story.append(KeepTogether([
        Paragraph("SCORE CLASSIFICATION GUIDE", S['section']),
        _build_legend_table(S, W),
        Spacer(1,8)
    ]))

    # ── Charts ──
    for key, title, h_ratio in [
        ("profile",     "IQ AND FACTOR SCORE PROFILE CHART", 0.52),
        ("gauge",       "FULL SCALE IQ CLASSIFICATION GAUGE", 0.38),
        ("discrepancy", "NV vs. V IQ DISCREPANCY ANALYSIS",  0.42),
        ("radar",       "FACTOR INDEX RADAR PROFILE",         0.55),
        ("subtest",     "SUBTEST COMPARISON CHART",           0.46),
    ]:
        if charts.get(key):
            img = RLImage(io.BytesIO(charts[key]), width=W, height=W*h_ratio)
            img.hAlign = 'CENTER'
            story.append(KeepTogether([
                Paragraph(title, S['section']),
                img,
                Spacer(1,8)
            ]))

    # ── Clinical Narrative ──
    story.append(HRFlowable(width=W, thickness=1, color=PDF_GOLD, spaceAfter=6))
    story.append(Paragraph("CLINICAL NARRATIVE REPORT", S['section']))

    sec_pat = re.compile(r'^\d+\.\s+[A-Z][A-Z\s,&/\(\):\']+$')
    header_words = {
        "REASON FOR REFERRAL","BACKGROUND INFORMATION","TESTS ADMINISTERED",
        "BEHAVIORAL OBSERVATIONS","ASSESSMENT RESULTS AND INTERPRETATION",
        "STRENGTHS AND WEAKNESSES PROFILE","DIAGNOSTIC IMPRESSIONS",
        "RECOMMENDATIONS","SUMMARY","PARENT-FRIENDLY SUMMARY",
        "STANFORD-BINET INTELLIGENCE SCALES, FIFTH EDITION — PSYCHOLOGICAL REPORT",
        "CLINICAL NARRATIVE REPORT",
    }
    in_parent = False

    for line in report_text.split('\n'):
        ls = line.strip()
        if not ls:
            story.append(Spacer(1,4)); continue
        if ls.startswith('━') or ls.startswith('═') or ls.startswith('---'):
            story.append(HRFlowable(width=W,thickness=0.4,color=PDF_BORDER,spaceAfter=4)); continue

        upper = ls.upper()
        is_sec = (sec_pat.match(ls) or ls in header_words or
                  any(upper.startswith(h) for h in header_words))

        if is_sec:
            if "PARENT" in upper or "SIMPLIFIED" in upper:
                in_parent = True
                story.append(HRFlowable(width=W,thickness=2,color=PDF_GOLD,spaceAfter=4))
                story.append(Paragraph(f"⭐  {ls}  ⭐", ParagraphStyle('psec',
                    fontName='Helvetica-Bold',fontSize=12,textColor=PDF_GOLD,
                    spaceBefore=10,spaceAfter=6,alignment=TA_CENTER)))
            else:
                story.append(Paragraph(ls, S['section']))
            continue

        if '|' in ls:
            parts = [p.strip() for p in ls.split('|') if p.strip()]
            if len(parts) >= 2:
                skip = [("name","date of birth"),("field","value")]
                key_pair = (parts[0].lower().strip(),parts[1].lower().strip()) if len(parts)>=2 else ("","")
                if key_pair not in skip:
                    col_w = W/len(parts)
                    mini = Table([[Paragraph(p,S['body']) for p in parts]], colWidths=[col_w]*len(parts))
                    mini.setStyle(TableStyle([
                        ('BOX',(0,0),(-1,-1),0.3,PDF_BORDER),
                        ('BACKGROUND',(0,0),(-1,-1),_shade1()),
                        ('TOPPADDING',(0,0),(-1,-1),3),
                        ('BOTTOMPADDING',(0,0),(-1,-1),3),
                        ('LEFTPADDING',(0,0),(-1,-1),5),
                    ]))
                    story.append(KeepTogether([mini])); continue

        style = S['parent'] if in_parent else S['body']
        story.append(Paragraph(ls, style))

    # ── Footer note ──
    story.append(Spacer(1,12))
    story.append(HRFlowable(width=W,thickness=0.5,color=PDF_BORDER))
    story.append(Spacer(1,4))
    footer_center = f" | {center_name}" if center_name else ""
    story.append(Paragraph(
        f"This report is strictly confidential{footer_center}. Results reflect performance at the time of testing only. "
        "Interpretation should be made in conjunction with clinical observation, history, and other assessment data. "
        "No formal diagnosis is provided herein.",
        S['small']))

    doc.build(story)
    buf.seek(0)
    return buf

def _build_score_table(data, S, W):
    rows = [[
        Paragraph('<b>Scale</b>',S['small']),
        Paragraph('<b>SS</b>',S['small']),
        Paragraph('<b>%ile</b>',S['small']),
        Paragraph('<b>90% CI</b>',S['small']),
        Paragraph('<b>Classification</b>',S['small']),
        Paragraph('<b>Band</b>',S['small']),
    ]]
    extra_cmds = []
    row_i = 1
    for k in IQ_SCORES:
        v = data.get(k)
        if not v: continue
        en_c,_,col = classify(v); pct=percentile_from_ss(v); ci=data.get(f"{k}_ci","—") or "—"
        rows.append([
            Paragraph(f"<b>{IQ_LABELS[k]['en']}</b>",S['body']),
            Paragraph(f"<b>{v}</b>",S['body']),
            Paragraph(f"{pct}th",S['body']),
            Paragraph(ci,S['body']),
            Paragraph(en_c,S['body']),
            Paragraph("",S['body']),
        ])
        extra_cmds.append(('BACKGROUND',(5,row_i),(5,row_i),colors.HexColor(col)))
        extra_cmds.append(('TEXTCOLOR', (5,row_i),(5,row_i),colors.black))
        extra_cmds.append(('FONTNAME',  (5,row_i),(5,row_i),'Helvetica-Bold'))
        row_i += 1
    # sub-header
    rows.append([Paragraph('<b>FACTOR INDEX SCORES</b>',S['small']),
                 Paragraph('',S['small']),Paragraph('',S['small']),
                 Paragraph('',S['small']),Paragraph('',S['small']),Paragraph('',S['small'])])
    extra_cmds.append(('BACKGROUND',(0,row_i),(-1,row_i),colors.HexColor('#2E6DB4')))
    extra_cmds.append(('TEXTCOLOR', (0,row_i),(-1,row_i),colors.white))
    extra_cmds.append(('FONTNAME',  (0,row_i),(-1,row_i),'Helvetica-Bold'))
    row_i += 1
    for k in FACTOR_SCORES:
        v = data.get(k)
        if not v: continue
        en_c,_,col = classify(v); pct=percentile_from_ss(v); ci=data.get(f"{k}_ci","—") or "—"
        bg = _shade2() if row_i%2==0 else colors.white
        rows.append([
            Paragraph(FACTOR_LABELS[k]['en'],S['body']),
            Paragraph(f"<b>{v}</b>",S['body']),
            Paragraph(f"{pct}th",S['body']),
            Paragraph(ci,S['body']),
            Paragraph(en_c,S['body']),
            Paragraph("",S['body']),
        ])
        extra_cmds.append(('BACKGROUND',(0,row_i),(4,row_i),bg))
        extra_cmds.append(('BACKGROUND',(5,row_i),(5,row_i),colors.HexColor(col)))
        extra_cmds.append(('TEXTCOLOR', (5,row_i),(5,row_i),colors.black))
        extra_cmds.append(('FONTNAME',  (5,row_i),(5,row_i),'Helvetica-Bold'))
        row_i += 1
    tbl = _score_tbl(rows,[6*cm,1.5*cm,1.8*cm,2.5*cm,4*cm,1.7*cm],extra_cmds)
    return tbl

def _build_subtest_table(data, S, W):
    has_any = any(data.get(f"nv_{k.lower()}") or data.get(f"v_{k.lower()}") for k in FACTOR_SCORES)
    if not has_any: return None
    rows = [[
        Paragraph('<b>Subtest</b>',S['small']),
        Paragraph('<b>Nonverbal SS</b>',S['small']),
        Paragraph('<b>Verbal SS</b>',S['small']),
        Paragraph('<b>NV Classification</b>',S['small']),
        Paragraph('<b>V Classification</b>',S['small']),
    ]]
    extra = []
    for i,k in enumerate(FACTOR_SCORES):
        nv=data.get(f"nv_{k.lower()}"); v=data.get(f"v_{k.lower()}")
        bg = _shade1() if i%2==0 else colors.white
        nv_c = classify(nv)[0] if nv else "—"
        v_c  = classify(v)[0]  if v  else "—"
        rows.append([
            Paragraph(FACTOR_LABELS[k]['en'],S['body']),
            Paragraph(f"<b>{nv}</b>" if nv else "—",S['body']),
            Paragraph(f"<b>{v}</b>"  if v  else "—",S['body']),
            Paragraph(nv_c,S['body']),
            Paragraph(v_c, S['body']),
        ])
        extra.append(('BACKGROUND',(0,i+1),(-1,i+1),bg))
    tbl = _score_tbl(rows,[5.5*cm,2.5*cm,2.5*cm,3.5*cm,3.5*cm],extra)
    return tbl

def _build_legend_table(S, W):
    entries = [
        ("#1565C0","T ≥ 130","Gifted / Very Gifted","Exceptional cognitive ability; highly enriched programming recommended."),
        ("#0288D1","120–129","Superior","Well above average; advanced academic enrichment beneficial."),
        ("#00897B","110–119","High Average","Above average; likely succeeds with standard instruction."),
        ("#388E3C","90–109","Average","Typical cognitive development for age."),
        ("#F9A825","80–89", "Low Average","Mild difficulties; monitoring and support may be needed."),
        ("#EF6C00","70–79", "Borderline","Significant difficulties; intervention recommended."),
        ("#E53935","55–69", "Mild Impairment","Substantial support and adapted instruction needed."),
        ("#B71C1C","40–54", "Moderate Impairment","Intensive support and specialized programming required."),
    ]
    rows = [[Paragraph('<b>Colour</b>',S['small']),Paragraph('<b>SS Range</b>',S['small']),
             Paragraph('<b>Classification</b>',S['small']),Paragraph('<b>Meaning</b>',S['small'])]]
    cmds = [
        ('BACKGROUND',(0,0),(-1,0),_hdr_style()),
        ('TEXTCOLOR', (0,0),(-1,0),colors.white),
        ('FONTNAME',  (0,0),(-1,0),'Helvetica-Bold'),
        ('FONTSIZE',  (0,0),(-1,-1),8.5),
        ('BOX',       (0,0),(-1,-1),0.5,PDF_BORDER),
        ('INNERGRID', (0,0),(-1,-1),0.3,PDF_BORDER),
        ('VALIGN',    (0,0),(-1,-1),'MIDDLE'),
        ('TOPPADDING',(0,0),(-1,-1),4),
        ('BOTTOMPADDING',(0,0),(-1,-1),4),
        ('LEFTPADDING',(0,0),(-1,-1),5),
    ]
    for i,(col,rng,classif,meaning) in enumerate(entries,start=1):
        rows.append([Paragraph("",S['small']),Paragraph(rng,S['body']),
                     Paragraph(f"<b>{classif}</b>",S['body']),Paragraph(meaning,S['body'])])
        cmds.append(('BACKGROUND',(0,i),(0,i),colors.HexColor(col)))
        cmds.append(('TEXTCOLOR', (0,i),(0,i),colors.black))
        cmds.append(('ROWBACKGROUNDS',(1,i),(-1,i),
                     [colors.white if i%2==0 else colors.HexColor('#F5F7FA')]))
    tbl = Table(rows, colWidths=[1.2*cm,2.5*cm,3.8*cm,W-7.5*cm], repeatRows=1)
    tbl.setStyle(TableStyle(cmds))
    return tbl

# ══════════════════════════════════════════════════════════════
#  STEP 4B: BUILD ARABIC WORD REPORT
# ══════════════════════════════════════════════════════════════
def build_word_doc(report_text: str, data: dict, charts: dict,
                   center_name: str = "", logo_bytes: bytes = None) -> io.BytesIO:
    is_rtl = True
    doc = Document()
    for sec_ in doc.sections:
        sec_.top_margin=Cm(2.0); sec_.bottom_margin=Cm(2.0)
        sec_.left_margin=Cm(2.5); sec_.right_margin=Cm(2.5)
    # Page border
    for sec_ in doc.sections:
        sp=sec_._sectPr; pb=OxmlElement('w:pgBorders')
        pb.set(qn('w:offsetFrom'),'page')
        for side in ('top','left','bottom','right'):
            b=OxmlElement(f'w:{side}'); b.set(qn('w:val'),'single')
            b.set(qn('w:sz'),'8'); b.set(qn('w:space'),'24')
            b.set(qn('w:color'),'1B3A6B'); pb.append(b)
        sp.append(pb)
    # Footer
    for sec_ in doc.sections:
        ft=sec_.footer; fp=ft.paragraphs[0] if ft.paragraphs else ft.add_paragraph()
        fp.clear(); fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r_=fp.add_run(); r_.font.size=Pt(9); r_.font.color.rgb=MID_BLUE_RGB
        for tag,text in [('begin',None),(None,' PAGE '),('end',None)]:
            if tag:
                el=OxmlElement('w:fldChar'); el.set(qn('w:fldCharType'),tag); r_._r.append(el)
            else:
                it=OxmlElement('w:instrText'); it.text=text; r_._r.append(it)

    def set_rtl(p):
        pPr=p._p.get_or_add_pPr(); pPr.append(OxmlElement("w:bidi"))
        jc=OxmlElement("w:jc"); jc.set(qn("w:val"),"right"); pPr.append(jc)

    def add_para(text,bold=False,size=11,color=None,space_before=0,space_after=4,alignment=None,italic=False):
        p=doc.add_paragraph()
        p.paragraph_format.space_before=Pt(space_before); p.paragraph_format.space_after=Pt(space_after)
        set_rtl(p)
        if alignment: p.alignment=alignment
        r_=p.add_run(text); r_.font.size=Pt(size); r_.font.name="Arial"
        r_.font.bold=bold; r_.font.italic=italic
        if color: r_.font.color.rgb=color
        return p

    def add_section_title(text):
        p=doc.add_paragraph()
        p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(4)
        p.paragraph_format.keep_with_next=True
        set_rtl(p)
        r_=p.add_run(text.strip()); r_.font.size=Pt(13); r_.font.name="Arial"
        r_.font.bold=True; r_.font.color.rgb=DEEP_BLUE_RGB
        pPr=p._p.get_or_add_pPr(); pBdr=OxmlElement('w:pBdr')
        bot=OxmlElement('w:bottom'); bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'8')
        bot.set(qn('w:space'),'2'); bot.set(qn('w:color'),'1B3A6B')
        pBdr.append(bot); pPr.append(pBdr)

    def make_table(col_widths):
        t=doc.add_table(rows=0,cols=len(col_widths)); t.style='Table Grid'
        try:
            tPr=t._tbl.tblPr
            bv=OxmlElement('w:bidiVisual'); tPr.append(bv)
            tW=OxmlElement('w:tblW'); tW.set(qn('w:w'),'9026'); tW.set(qn('w:type'),'dxa'); tPr.append(tW)
            tg=OxmlElement('w:tblGrid')
            for w in col_widths:
                gc=OxmlElement('w:gridCol'); gc.set(qn('w:w'),str(w)); tg.append(gc)
            t._tbl.insert(0,tg)
        except: pass
        return t

    def add_table_row(table,cells_data,is_header=False,shade=None):
        row=table.add_row()
        trPr=row._tr.get_or_add_trPr()
        cs=OxmlElement('w:cantSplit'); cs.set(qn('w:val'),'1'); trPr.append(cs)
        bidi_=OxmlElement('w:bidi'); trPr.append(bidi_)
        for cell,(txt,bold_) in zip(row.cells,cells_data):
            cell.text=""
            p=cell.paragraphs[0]
            pPr=p._p.get_or_add_pPr()
            pPr.append(OxmlElement("w:bidi"))
            jc=OxmlElement("w:jc"); jc.set(qn("w:val"),"right"); pPr.append(jc)
            vr=p.add_run(str(txt)); vr.font.size=Pt(9.5); vr.font.name="Arial"; vr.font.bold=bold_
            if is_header: vr.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
            else: vr.font.color.rgb=RGBColor(0,0,0)
            tc=cell._tc; tcP=tc.get_or_add_tcPr()
            shd=OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
            if is_header: shd.set(qn('w:fill'),'1B3A6B')
            elif shade: shd.set(qn('w:fill'),shade)
            else: shd.set(qn('w:fill'),'FFFFFF')
            tcP.append(shd)
            mg=OxmlElement('w:tcMar')
            for side in ['top','bottom','left','right']:
                m=OxmlElement(f'w:{side}'); m.set(qn('w:w'),'70'); m.set(qn('w:type'),'dxa'); mg.append(m)
            tcP.append(mg)

    # ── Header ──
    p_hdr=doc.add_paragraph(); p_hdr.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p_hdr.paragraph_format.space_after=Pt(6)

    # Logo: uploaded bytes take priority, fallback to file
    logo_added = False
    if logo_bytes:
        try:
            p_hdr.add_run().add_picture(io.BytesIO(logo_bytes), width=Inches(2.8))
            logo_added = True
        except: pass
    if not logo_added and os.path.exists(LOGO_PATH):
        try:
            p_hdr.add_run().add_picture(LOGO_PATH, width=Inches(2.8))
        except: pass

    # Center name
    if center_name:
        p_cn = doc.add_paragraph(); p_cn.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p_cn.paragraph_format.space_after=Pt(3)
        set_rtl(p_cn)
        r_cn = p_cn.add_run(center_name)
        r_cn.font.name="Arial"; r_cn.font.size=Pt(14)
        r_cn.font.bold=True; r_cn.font.color.rgb=DEEP_BLUE_RGB
    r_t=p_hdr.add_run("\nمقياس ستانفورد-بينيه للذكاء — الصورة الخامسة\nتقرير التقييم النفسي")
    r_t.font.name="Arial"; r_t.font.size=Pt(16); r_t.font.bold=True; r_t.font.color.rgb=DEEP_BLUE_RGB
    add_para("SB5 · ترجمة وتقنين أ.د/ صفوت فرج",size=9,color=GOLD_RGB,
             alignment=WD_ALIGN_PARAGRAPH.CENTER,space_after=2)
    # Gold rule
    p_sep=doc.add_paragraph(); p_sep.paragraph_format.space_before=Pt(2); p_sep.paragraph_format.space_after=Pt(10)
    pPr=p_sep._p.get_or_add_pPr(); pBdr2=OxmlElement('w:pBdr')
    bot2=OxmlElement('w:bottom'); bot2.set(qn('w:val'),'single'); bot2.set(qn('w:sz'),'12')
    bot2.set(qn('w:space'),'2'); bot2.set(qn('w:color'),'C8922A')
    pBdr2.append(bot2); pPr.append(pBdr2)

    # ── Demographics ──
    add_section_title("بيانات المفحوص")
    info_tbl=make_table([2200,3000,2200,3000])
    add_table_row(info_tbl,[("الحقل",True),("",True),("الحقل",True),("",True)],is_header=True)
    add_table_row(info_tbl,[(data.get("dob","—") or "—",False),("تاريخ الميلاد",True),(data.get("name","—") or "—",False),("الاسم",True)])
    add_table_row(info_tbl,[(data.get("gender","—") or "—",False),("النوع",True),(data.get("age","—") or "—",False),("العمر",True)],shade="EEF2FF")
    add_table_row(info_tbl,[(data.get("test_date","—") or "—",False),("تاريخ التطبيق",True),(data.get("examiner","—") or "—",False),("الفاحص",True)])
    add_table_row(info_tbl,[(date.today().strftime('%Y/%m/%d'),False),("تاريخ التقرير",True),(data.get("referral","—") or "—",False),("جهة الإحالة",True)],shade="EEF2FF")
    doc.add_paragraph().paragraph_format.space_after=Pt(6)

    # ── Score Summary Table ──
    add_section_title("ملخص درجات الذكاء ومؤشرات العوامل")
    score_tbl=make_table([3500,1200,1100,2200,2000])
    add_table_row(score_tbl,[("المقياس",True),("الدرجة",True),("المئيني",True),("مدى الثقة",True),("الفئة",True)],is_header=True)
    for k in IQ_SCORES:
        v=data.get(k)
        if not v: continue
        _,ar_c,_ = classify(v); pct=percentile_from_ss(v); ci=data.get(f"{k}_ci","—") or "—"
        shade="F5F7FF" if IQ_SCORES.index(k)%2==0 else "FFFFFF"
        add_table_row(score_tbl,[(IQ_LABELS[k]['ar'],True),(str(v),False),(f"{pct}%",False),(ci,False),(ar_c,False)],shade=shade)
    add_table_row(score_tbl,[("درجات مؤشرات العوامل",True),("",False),("",False),("",False),("",False)],is_header=True)
    for i,k in enumerate(FACTOR_SCORES):
        v=data.get(k)
        if not v: continue
        _,ar_c,_ = classify(v); pct=percentile_from_ss(v); ci=data.get(f"{k}_ci","—") or "—"
        shade="FFF8EE" if i%2==0 else "FFFFFF"
        add_table_row(score_tbl,[(FACTOR_LABELS[k]['ar'],True),(str(v),False),(f"{pct}%",False),(ci,False),(ar_c,False)],shade=shade)
    doc.add_paragraph().paragraph_format.space_after=Pt(6)

    # ── Subtest Table ──
    has_sub = any(data.get(f"nv_{k.lower()}") or data.get(f"v_{k.lower()}") for k in FACTOR_SCORES)
    if has_sub:
        add_section_title("الدرجات المعيارية للاختبارات الفرعية")
        sub_tbl=make_table([4000,2000,2000])
        add_table_row(sub_tbl,[("الاختبار الفرعي",True),("غير لفظي",True),("لفظي",True)],is_header=True)
        for i,k in enumerate(FACTOR_SCORES):
            nv=data.get(f"nv_{k.lower()}","—") or "—"; v=data.get(f"v_{k.lower()}","—") or "—"
            shade="F5F7FF" if i%2==0 else "FFFFFF"
            add_table_row(sub_tbl,[(FACTOR_LABELS[k]['ar'],True),(str(nv),False),(str(v),False)],shade=shade)
        doc.add_paragraph().paragraph_format.space_after=Pt(6)

    # ── Charts ──
    for chart_key,chart_title_ar in [
        ("profile",     "ملف الدرجات الكاملة"),
        ("gauge",       "مقياس تصنيف الذكاء الكلي"),
        ("discrepancy", "تحليل الفجوة بين الذكاء اللفظي وغير اللفظي"),
        ("radar",       "مخطط رادار مؤشرات العوامل"),
        ("subtest",     "مقارنة الاختبارات الفرعية اللفظية وغير اللفظية"),
    ]:
        if charts.get(chart_key):
            add_section_title(chart_title_ar)
            p_c=doc.add_paragraph(); p_c.alignment=WD_ALIGN_PARAGRAPH.CENTER
            p_c.paragraph_format.space_after=Pt(8)
            w=Inches(5.8) if chart_key in ("profile","subtest","discrepancy") else Inches(4.0)
            p_c.add_run().add_picture(io.BytesIO(charts[chart_key]),width=w)

    # ── Narrative ──
    add_section_title("التقرير السريري التفصيلي")
    sec_ar_pat=re.compile(r'^[١٢٣٤٥٦٧٨٩أبجدهوزحطيكلمنسعفصقرشت\d]+[\.،:]\s+[\u0600-\u06FF]')
    header_words_ar={"سبب الإحالة","المعلومات الأساسية","الاختبارات المطبقة","الملاحظات السلوكية",
                     "نتائج التقييم وتفسيرها","نقاط القوة والقصور","الانطباعات التشخيصية",
                     "التوصيات","الملخص","ملخص للوالدين"}
    in_parent_ar=False; in_table_=False; cur_tbl_=None

    for line in report_text.split('\n'):
        ls=line.strip()
        if not ls:
            if in_table_: in_table_=False; cur_tbl_=None
            doc.add_paragraph().paragraph_format.space_after=Pt(2); continue
        if ls.startswith('━') or ls.startswith('═') or ls.startswith('---'):
            in_table_=False; cur_tbl_=None; continue
        upper=ls.upper()
        is_sec=(sec_ar_pat.match(ls) or any(ls.startswith(h) or upper.startswith(h.upper()) for h in header_words_ar))
        if is_sec:
            in_table_=False; cur_tbl_=None
            if "والدين" in ls or "مبسط" in ls:
                in_parent_ar=True
                p_sep2=doc.add_paragraph()
                pPr2=p_sep2._p.get_or_add_pPr(); pBdr3=OxmlElement('w:pBdr')
                t2=OxmlElement('w:top'); t2.set(qn('w:val'),'single'); t2.set(qn('w:sz'),'8')
                t2.set(qn('w:space'),'4'); t2.set(qn('w:color'),'C8922A')
                pBdr3.append(t2); pPr2.append(pBdr3)
                p2=doc.add_paragraph(); set_rtl(p2)
                r2_=p2.add_run("⭐ "+ls+" ⭐")
                r2_.font.size=Pt(13); r2_.font.name="Arial"; r2_.font.bold=True; r2_.font.color.rgb=GOLD_RGB
            else:
                add_section_title(ls)
            continue
        if '|' in ls:
            parts=[p.strip() for p in ls.split('|') if p.strip()]
            if not parts: continue
            if not in_table_ or cur_tbl_ is None:
                in_table_=True; w_ea=9026//len(parts); cur_tbl_=make_table([w_ea]*len(parts))
            add_table_row(cur_tbl_,[(p,False) for p in parts],shade="F5F7FF"); continue
        in_table_=False; cur_tbl_=None
        sz=11 if in_parent_ar else 10.5
        col_=DARK_RGB if in_parent_ar else None
        add_para(ls,size=sz,space_after=3,color=col_)

    doc.add_paragraph().paragraph_format.space_after=Pt(12)
    add_para("هذا التقرير سري. النتائج تعكس الأداء في تاريخ التطبيق فقط. يجب تفسير النتائج في سياق الملاحظة الإكلينيكية والبيانات الأخرى.",
             size=8,color=MID_BLUE_RGB,italic=True)
    buf=io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf

# ══════════════════════════════════════════════════════════════
#  EMAIL
# ══════════════════════════════════════════════════════════════
def send_email(data, buf_pdf, buf_doc, fn_pdf, fn_doc):
    name=data.get("name","—") or "—"
    fsiq=data.get("FSIQ")
    center=data.get("_center_name","") or ""
    en_c,ar_c,_=classify(fsiq) if fsiq else ("—","—","#888")
    date_str=date.today().strftime('%B %d, %Y')
    center_line = f"<tr><td style='padding:5px 0;color:#555;width:40%;'>Center</td><td><strong>{center}</strong></td></tr>" if center else ""
    msg=MIMEMultipart('mixed')
    msg['From']=GMAIL_USER; msg['To']=RECIPIENT_EMAIL
    subject_center = f" | {center}" if center else ""
    msg['Subject']=f"[SB5 Report] {name}{subject_center} — {date_str}"
    body=f"""<html><body style="font-family:Georgia,serif;color:#1A1A2E;background:#F5F7FA;padding:20px;">
  <div style="max-width:560px;margin:0 auto;background:white;border:1px solid #C8922A;border-radius:8px;padding:28px;">
    <h2 style="font-weight:600;font-size:18px;color:#1B3A6B;margin-bottom:2px;">Stanford-Binet 5 — Assessment Report</h2>
    <p style="color:#888;font-size:11px;margin-top:0;">SB5 Psychological Assessment — Auto-generated from uploaded report</p>
    <hr style="border:none;border-top:2px solid #C8922A;margin:16px 0;">
    <table style="width:100%;font-size:13px;border-collapse:collapse;">
      {center_line}
      <tr><td style="padding:5px 0;color:#555;width:40%;">Examinee</td><td><strong>{name}</strong></td></tr>
      <tr><td style="padding:5px 0;color:#555;">Age</td><td>{data.get("age","—") or "—"}</td></tr>
      <tr><td style="padding:5px 0;color:#555;">FSIQ</td><td><strong style="color:#1B3A6B;">{fsiq or "—"}</strong></td></tr>
      <tr><td style="padding:5px 0;color:#555;">Classification</td><td>{en_c} / {ar_c}</td></tr>
      <tr><td style="padding:5px 0;color:#555;">Examiner</td><td>{data.get("examiner","—") or "—"}</td></tr>
      <tr><td style="padding:5px 0;color:#555;">Report Date</td><td>{date_str}</td></tr>
    </table>
    <hr style="border:none;border-top:1px solid #DDE5F8;margin:16px 0;">
    <p style="font-size:12px;line-height:1.7;">Two reports attached:<br>
    📄 <strong>English Report (PDF)</strong> — Premium clinical report with charts and tables<br>
    📝 <strong>Arabic Report (Word)</strong> — تقرير سريري عربي مُحسَّن بالجداول والمخططات</p>
    <p style="font-size:10px;color:#888;font-style:italic;">Confidential — for the evaluating clinician only.</p>
  </div></body></html>"""
    msg.attach(MIMEText(body,'html'))
    # PDF
    buf_pdf.seek(0)
    part_pdf=MIMEBase('application','pdf')
    part_pdf.set_payload(buf_pdf.read()); encoders.encode_base64(part_pdf)
    part_pdf.add_header('Content-Disposition','attachment',filename=fn_pdf)
    msg.attach(part_pdf)
    # Word
    buf_doc.seek(0)
    part_doc=MIMEBase('application','vnd.openxmlformats-officedocument.wordprocessingml.document')
    part_doc.set_payload(buf_doc.read()); encoders.encode_base64(part_doc)
    part_doc.add_header('Content-Disposition','attachment',filename=fn_doc)
    msg.attach(part_doc)
    with smtplib.SMTP_SSL('smtp.gmail.com',465) as srv:
        srv.login(GMAIL_USER,GMAIL_PASS)
        srv.sendmail(GMAIL_USER,RECIPIENT_EMAIL,msg.as_string())

# ══════════════════════════════════════════════════════════════
#  PAGE CONFIG & CSS
# ══════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="Stanford-Binet 5 Report",
    page_icon="🎓",
    layout="centered",
    initial_sidebar_state="collapsed",
)
st.markdown(f"""
<style>
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:wght@400;700&family=Inter:wght@300;400;500;600&display=swap');
html,body,[class*="css"]{{font-family:'Inter',sans-serif;background:{LIGHT_BG};}}
.stApp{{background:{LIGHT_BG};}}
#MainMenu{{visibility:hidden!important;}}
header[data-testid="stHeader"]{{visibility:hidden!important;}}
footer{{visibility:hidden!important;}}
[data-testid="stToolbar"]{{display:none!important;}}

.sb5-header{{
    background:linear-gradient(135deg,{DEEP_BLUE} 0%,{MID_BLUE} 60%,{DEEP_BLUE} 100%);
    border-radius:16px;padding:28px 36px;margin-bottom:24px;
    box-shadow:0 6px 24px rgba(27,58,107,0.3);border-bottom:4px solid {GOLD};
    text-align:center;
}}
.sb5-header h1{{color:white;font-family:'Playfair Display',serif;font-size:1.9rem;font-weight:700;margin:0 0 6px;}}
.sb5-header p{{color:#A8BFDD;font-size:0.87rem;margin:0;}}

.upload-card{{
    background:white;border-radius:12px;padding:32px 36px;
    box-shadow:0 2px 16px rgba(27,58,107,0.1);
    border:2px dashed {MID_BLUE};text-align:center;margin:16px 0;
}}
.upload-card h3{{color:{DEEP_BLUE};font-size:1.2rem;margin-bottom:8px;}}
.upload-card p{{color:#666;font-size:0.88rem;line-height:1.7;}}

.thank-you{{text-align:center;padding:4rem 2rem;}}
.thank-you h2{{font-family:'Playfair Display',serif;font-size:2rem;font-weight:400;color:{DEEP_BLUE};margin-bottom:1rem;}}
.thank-you p{{color:{MID_BLUE};font-size:.95rem;line-height:1.9;}}

.stButton>button{{
    background:{DEEP_BLUE}!important;color:white!important;border:none!important;
    border-radius:10px!important;padding:12px 32px!important;font-size:14px!important;
    font-weight:600!important;box-shadow:0 3px 12px rgba(27,58,107,0.3)!important;
    transition:all 0.2s!important;
}}
.stButton>button:hover{{background:{MID_BLUE}!important;}}
</style>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════
#  HEADER
# ══════════════════════════════════════════════════════════════
col_logo, col_hdr = st.columns([1,5])
with col_logo:
    if os.path.exists(LOGO_PATH):
        st.image(LOGO_PATH, width=110)
with col_hdr:
    st.markdown("""
    <div class="sb5-header">
        <h1>🎓 Stanford-Binet Intelligence Scales, 5th Ed.</h1>
        <p>مقياس ستانفورد-بينيه للذكاء — الصورة الخامسة&nbsp;&nbsp;·&nbsp;&nbsp;
        Upload your Arabic SB5 report → receive premium English PDF + Arabic Word reports by email</p>
    </div>""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════
#  SESSION STATE
# ══════════════════════════════════════════════════════════════
if "done" not in st.session_state: st.session_state.done = False

# ══════════════════════════════════════════════════════════════
#  THANK-YOU SCREEN
# ══════════════════════════════════════════════════════════════
if st.session_state.done:
    data = st.session_state.get("last_data",{})
    name = data.get("name","—") or "—"
    fsiq = data.get("FSIQ")
    center_done = data.get("_center_name","")
    en_c,_,_ = classify(fsiq) if fsiq else ("—","—","#888")
    # Show uploaded logo if available, else file logo
    logo_b = st.session_state.get("center_logo_bytes")
    if logo_b:
        c1,c2,c3=st.columns([1,2,1])
        with c2: st.image(logo_b, use_container_width=True)
    elif os.path.exists(LOGO_PATH):
        c1,c2,c3=st.columns([1,2,1])
        with c2: st.image(LOGO_PATH,use_container_width=True)
    center_line = f"<p style='color:{MID_BLUE};font-size:.9rem;margin-bottom:.5rem;'><strong>{center_done}</strong></p>" if center_done else ""
    st.markdown(f"""<div class="thank-you">
        {center_line}
        <h2>Reports Sent Successfully</h2>
        <p><strong>{name}</strong></p>
        <p>FSIQ: <strong>{fsiq or "—"}</strong> &nbsp;·&nbsp; {en_c}</p>
        <p style="margin-top:1.2rem;font-size:.87rem;">
            The English PDF and Arabic Word reports have been sent to the clinic email.<br>
            تم إرسال التقرير الإنجليزي (PDF) والتقرير العربي (Word) إلى البريد الإلكتروني للعيادة.
        </p>
    </div>""", unsafe_allow_html=True)
    _,btn_col,_ = st.columns([2,2,2])
    with btn_col:
        if st.button("↺ Upload New Report", use_container_width=True):
            for k in list(st.session_state.keys()): del st.session_state[k]
            st.rerun()
    st.stop()

# ══════════════════════════════════════════════════════════════
#  UPLOAD SECTION
# ══════════════════════════════════════════════════════════════
st.markdown("""<div class="upload-card">
    <h3>Upload Arabic SB5 Report</h3>
    <p>Upload the Arabic <strong>Word (.docx)</strong> report generated by your offline SB5 software.<br>
    The app will automatically extract all scores and information, generate a premium English PDF report
    and an enhanced Arabic Word report, and send both to the clinic email.<br><br>
    <strong>الرجاء رفع التقرير العربي بصيغة Word (.docx)</strong> — سيتم استخراج البيانات تلقائياً وإرسال التقريرين.</p>
</div>""", unsafe_allow_html=True)


# ── Center info ──
st.markdown("<div style='background:white;border-radius:10px;padding:20px 24px;"
            "box-shadow:0 2px 12px rgba(27,58,107,0.08);border-left:4px solid #C8922A;margin-bottom:16px;'>",
            unsafe_allow_html=True)
col_cn, col_logo_up = st.columns([3, 2])
with col_cn:
    st.markdown("<div style='font-size:.75rem;font-weight:700;color:#1B3A6B;"
                "text-transform:uppercase;letter-spacing:.08em;margin-bottom:6px;'>"
                "Center / Clinic Name &nbsp;·&nbsp; اسم المركز / العيادة</div>",
                unsafe_allow_html=True)
    center_name = st.text_input("center_name", value="",
                                placeholder="e.g. Wijdan Therapy Center / مركز وجدان للعلاج النفسي",
                                label_visibility="collapsed", key="center_name_inp")
with col_logo_up:
    st.markdown("<div style='font-size:.75rem;font-weight:700;color:#1B3A6B;"
                "text-transform:uppercase;letter-spacing:.08em;margin-bottom:6px;'>"
                "Center Logo (optional)</div>", unsafe_allow_html=True)
    logo_upload = st.file_uploader("Center logo", type=["png","jpg","jpeg"],
                                   label_visibility="collapsed", key="logo_upload_inp")
    if logo_upload:
        st.session_state["center_logo_bytes"] = logo_upload.read()
        st.success("✓ Logo uploaded")
st.markdown("</div>", unsafe_allow_html=True)

st.markdown("<div style='font-size:.75rem;font-weight:700;color:#1B3A6B;"
            "text-transform:uppercase;letter-spacing:.08em;margin-bottom:6px;'>"
            "Arabic SB5 Report (.docx)</div>", unsafe_allow_html=True)

uploaded = st.file_uploader(
    "Choose Arabic SB5 Word report",
    type=["docx"],
    label_visibility="collapsed",
)

if uploaded:
    safe_name = re.sub(r'[^\w\-]','_', uploaded.name.replace('.docx',''))
    fn_pdf = f"{safe_name}_EN_Report.pdf"
    fn_doc = f"{safe_name}_AR_Report.docx"

    with st.spinner("⏳ Reading report, extracting data, generating charts and reports — please wait 30–60 seconds..."):

        # 1. Extract raw text
        raw_text = extract_docx_text(uploaded)
        if not raw_text.strip():
            st.error("Could not extract text from this Word document. Please make sure it is a valid .docx file.")
            st.stop()

        # 2. Parse structured data
        data = extract_data_from_text(raw_text)
        if not data:
            st.error("Could not parse the report data. Please check the uploaded file.")
            st.stop()

        # 3. Generate charts
        charts = {}
        pb = make_profile_chart(data)
        if pb: charts["profile"] = pb

        fsiq_val = data.get("FSIQ")
        if fsiq_val:
            charts["gauge"] = make_classification_gauge(fsiq_val)

        disc_b = make_discrepancy_chart(data)
        if disc_b: charts["discrepancy"] = disc_b

        factor_ss = {k: data.get(k) for k in FACTOR_SCORES}
        radar_b = make_factor_radar(data)
        if radar_b: charts["radar"] = radar_b

        sub_b = make_subtest_chart(data)
        if sub_b: charts["subtest"] = sub_b

        # 4. Generate narrative reports
        report_en = generate_en_report(data)
        report_ar = generate_ar_report(data)

        # 5. Build files
        center_name_v = center_name.strip() if center_name else ""
        logo_bytes_v  = st.session_state.get("center_logo_bytes", None)
        buf_pdf = build_pdf_report(report_en, data, charts, center_name_v, logo_bytes_v)
        buf_doc = build_word_doc(report_ar, data, charts, center_name_v, logo_bytes_v)

        # 6. Send email
        try:
            send_email(data, buf_pdf, buf_doc, fn_pdf, fn_doc)
        except Exception as e:
            st.warning(f"Report generated but email failed: {e}")

        # 7. Save state and redirect
        data["_center_name"] = center_name_v
        st.session_state["last_data"] = data
        st.session_state.done = True
        st.rerun()
